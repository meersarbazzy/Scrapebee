import os
import requests
import uuid
import pandas as pd
from bs4 import BeautifulSoup, Comment
from urllib.parse import urljoin, urlparse
from docx import Document
from pathvalidate import sanitize_filename
import urllib3
import time

# Suppress SSL warnings
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

class TIPPScraper:
    def __init__(self, output_dir="c:/Project-DQ/Validata/Scrapper/universal_output_v2"):
        self.base_url = "https://tipp.gov.pk/"
        self.output_dir = output_dir
        self.metadata_file = os.path.join(output_dir, "metadata.xlsx")
        self.metadata = [] # List to store dicts for DataFrame
        
        self.session = requests.Session()
        self.session.headers.update({
             'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def fetch(self, url):
        try:
            response = self.session.get(url, verify=False, timeout=20)
            if response.status_code == 500:
                print(f"WARNING: 500 Error for {url}, but attempting to parse content.")
                return response.text
            response.raise_for_status()
            return response.text
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 500:
                 return e.response.text
            print(f"Error fetching {url}: {e}")
            return None
        except Exception as e:
            print(f"Error fetching {url}: {e}")
            return None

    def run(self):
        print("Starting TIPP Scraper V2...")
        html = self.fetch(self.base_url)
        if not html:
            print("Failed to fetch homepage. Exiting.")
            return

        soup = BeautifulSoup(html, 'html.parser')
        
        # Parse Main Navigation
        # Heuristic: Find the main menu container. 
        # TIPP likely uses standard nav or ul structure.
        # We will look for deep nested ULs.
        
        # Try to find the nav structure
        # Common selectors: #menu-main-menu, .navbar-nav, etc.
        # Based on snippet from previous run, it's HTML5.
        
        nav_container = soup.find('nav') or soup.find(id='primary-menu') or soup.find(class_='navbar-nav')
        
        if not nav_container:
            # Fallback: Find the UL with the most links
            uls = soup.find_all('ul')
            if uls:
                # heuristic: pick the one with decent number of children, near top
                nav_container = uls[0] 
                # This is risky, but better than nothing for a blind run.
                # Improvements can be made after first run inspection.

        if nav_container:
            # Locate top-level LIs
             # If nav_container is 'nav', find 'ul' inside first
            if nav_container.name == 'nav':
                ul = nav_container.find('ul')
                if ul:
                    self.recursive_traverse(ul, [])
            elif nav_container.name == 'ul':
                self.recursive_traverse(nav_container, [])
            else:
                # It's a div maybe?
                ul = nav_container.find('ul')
                if ul:
                     self.recursive_traverse(ul, [])
        else:
             print("Could not locate Main Menu. Please refine selector.")

        # Save Excel at the end
        self.save_excel()
        print("Scraping Completed. Metadata saved.")

    def recursive_traverse(self, ul_element, path):
        # Iterate over direct children li
        children = ul_element.find_all('li', recursive=False)
        
        for li in children:
            # Extract Text and Link
            a_tag = li.find('a', recursive=False)
            if not a_tag:
                 # Check if 'a' is wrapped in div/span?
                 a_tag = li.find('a')
            
            if not a_tag:
                continue

            text = a_tag.get_text(strip=True)
            href = a_tag.get('href')
            
            # Submenu Check
            submenu = li.find('ul')
            
            new_path = path + [text]
            
            if submenu:
                # It's a Category
                # Some menus have the parent as a clickable link too.
                # Recursion first? Or Visit first?
                # Usually recursion.
                print(f"Traversing Category: {' > '.join(new_path)}")
                self.recursive_traverse(submenu, new_path)
            else:
                # It's a Leaf Page
                if href and href != '#' and not href.startswith('javascript'):
                    full_url = urljoin(self.base_url, href)
                    # Filter external links? 
                    if self.base_url in full_url:
                        print(f"  Scraping Page: {text} ({full_url})")
                        self.visit_and_extract(full_url, new_path)

    def visit_and_extract(self, url, path):
        html = self.fetch(url)
        if not html:
            return

        soup = BeautifulSoup(html, 'html.parser')
        
        # Metadata Prep
        main_cat = path[0] if len(path) > 0 else "Uncategorized"
        sub_cat = path[1] if len(path) > 1 else "General"
        
        # Content Detection Logic
        # Look for headers
        main_content = soup.find('main') or soup.find(id='content') or soup.body
        
        if not main_content:
            return

        # Removal of script/style
        for element in main_content(['script', 'style', 'nav', 'footer', 'iframe']):
            element.decompose()

        # Block Extraction Strategy
        # Iterate through headers H1-H6
        # Or generic containers if headers are missing
        
        blocks_found = False
        headers = main_content.find_all(['h1', 'h2', 'h3', 'h4'])
        
        for h in headers:
            title = h.get_text(strip=True)
            if not title:
                continue

            # Gather content
            body_text = []
            ext_url = None
            
            curr = h.next_sibling
            while curr:
                if curr.name in ['h1', 'h2', 'h3', 'h4', 'footer', 'nav']:
                    break
                
                if curr.name == 'a':
                    if not ext_url: # grab first link? or specific one?
                         ext_url = curr.get('href')
                elif curr.name:
                     # Check for links inside nested elements
                     if not ext_url:
                         child_a = curr.find('a')
                         if child_a:
                             ext_url = child_a.get('href')
                             
                     text = curr.get_text(strip=True)
                     if text:
                         body_text.append(text)
                
                curr = curr.next_sibling
            
            full_body = "\n".join(body_text)
            
            if full_body or ext_url:
                self.save_record(
                    title=title,
                    body=full_body,
                    ext_href=ext_url,
                    path=path,
                    source_url=url,
                    main_cat=main_cat,
                    sub_cat=sub_cat
                )
                blocks_found = True

        if not blocks_found:
             # Fallback: Treat whole page as one block
             # Use Page Title
             page_title = soup.title.get_text(strip=True) if soup.title else path[-1]
             # Extract all text
             full_body = main_content.get_text(separator='\n', strip=True)
             
             # Extract first meaningful link if any
             first_link = main_content.find('a', href=True)
             ext_url = first_link['href'] if first_link else None
             
             self.save_record(
                title=page_title,
                body=full_body[:5000], # Limit body size just in case?
                ext_href=ext_url,
                path=path,
                source_url=url,
                main_cat=main_cat,
                sub_cat=sub_cat
            )

    def save_record(self, title, body, ext_href, path, source_url, main_cat, sub_cat):
        # 1. UUID
        record_uuid = str(uuid.uuid4())
        
        # 2. Normalize External URL
        final_ext_url = None
        if ext_href:
            if ext_href.startswith('http'):
                final_ext_url = ext_href
            elif ext_href.startswith('/'):
                final_ext_url = urljoin(self.base_url, ext_href)
            # ignore javascript/hash
            if final_ext_url and ('javascript' in final_ext_url.lower() or final_ext_url == '#'):
                final_ext_url = None

        # 3. Create Files
        # Folder Structure
        folder_names = [sanitize_filename(p) for p in path] # Breadcrumbs folders
        save_path = os.path.join(self.output_dir, *folder_names)
        
        if not os.path.exists(save_path):
            os.makedirs(save_path)

        # File Name
        safe_title = sanitize_filename(title)
        if not safe_title:
             safe_title = "Untitled"
        filename = f"{safe_title[:50]}.docx"
        file_full_path = os.path.join(save_path, filename)
        
        # Avoid overwriting? Or append UUID?
        if os.path.exists(file_full_path):
             filename = f"{safe_title[:40]}_{record_uuid[:8]}.docx"
             file_full_path = os.path.join(save_path, filename)

        try:
            # Generate Word Doc
            doc = Document()
            
            # Custom Footer Hack (python-docx handles footers via sections)
            section = doc.sections[0]
            footer = section.footer
            p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            
            footer_text = ""
            if final_ext_url:
                footer_text = f"Official Reference Link: {final_ext_url}"
            else:
                 pass # User said "If no link exists... leave Excel cell EMPTY". 
                      # For footer, logic says "Strictly format...". 
                      # If empty, maybe omit text or put "Official Reference Link: N/A"?
                      # Implied: If no link, maybe no footer text needed? 
                      # We'll skip adding text if None.
            
            p_footer.text = footer_text

            # Header info in Body
            # [Main Category] > [Sub Category] > [Page Title]
            breadcrumb_text = f"{main_cat} > {sub_cat} > {title}"
            p_head = doc.add_paragraph()
            runner = p_head.add_run(breadcrumb_text)
            runner.bold = True
            
            doc.add_heading(title, level=1)
            doc.add_paragraph(body)
            
            doc.save(file_full_path)
            
            # 4. Add to Metadata
            relative_path = os.path.relpath(file_full_path, start=self.output_dir)
            
            self.metadata.append({
                "UUID": record_uuid,
                "Main_Category": main_cat,
                "Sub_Category": sub_cat,
                "Entity_Title": title, # Original text
                "Description_Snippet": body[:150],
                "Local_File_Path": relative_path,
                "TIPP_Source_URL": source_url,
                "External_Reference_URL": final_ext_url if final_ext_url else None # Blank/NaN
            })
            
            print(f"    Saved Record: {title} -> {filename}")
            
        except Exception as e:
            print(f"Error saving record {title}: {e}")

    def save_excel(self):
        if not self.metadata:
            print("No metadata to save.")
            return

        df = pd.DataFrame(self.metadata)
        # Enforce columns order
        cols = [
            "UUID", "Main_Category", "Sub_Category", "Entity_Title", 
            "Description_Snippet", "Local_File_Path", "TIPP_Source_URL", 
            "External_Reference_URL"
        ]
        
        # ensure exists (in case missing)
        for c in cols:
            if c not in df.columns:
                df[c] = None
                
        df = df[cols]
        df.to_excel(self.metadata_file, index=False)
        print(f"Master Index Saved: {self.metadata_file}")

if __name__ == "__main__":
    scraper = TIPPScraper()
    scraper.run()
