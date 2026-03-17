import os
import time
import uuid
import re
import pandas as pd
from typing import List, Dict, Any, Set
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathvalidate import sanitize_filename
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from selenium.common.exceptions import TimeoutException
from scrapebee.core.pdf_processor import download_and_clean_pdf

class GenericModScraper:
    def __init__(self, start_url, max_pages=50, output_dir=None, progress_callback=None):
        self.start_url = start_url
        self.base_domain = urlparse(start_url).netloc
        self.max_pages = max_pages
        # Dynamic path management: default to project root / data / outputs
        self.output_dir = output_dir or os.path.join(os.getcwd(), "data", "outputs")
        self.progress_callback = progress_callback
        self.metadata_file = os.path.join(self.output_dir, "metadata.xlsx")
        self.metadata = []
        self.visited_urls = set()
        self.crawl_queue = []
        
        # Setup Driver
        options = webdriver.ChromeOptions()
        options.add_argument("--start-maximized")
        options.add_argument("--page-load-strategy=eager") 
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        options.add_argument('--blink-settings=imagesEnabled=false')
        options.add_argument("--ignore-certificate-errors")
        options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        self.driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=options)
        
        # Anti-detection CDP
        self.driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
            "source": "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        })
        self.driver.set_page_load_timeout(60)
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir, exist_ok=True)

    def run(self):
        try:
            print(f"Starting Generic Crawl on: {self.start_url}")
            self.crawl_queue = [self.start_url]
            self.visited_urls.add(self.start_url)
            
            pages_processed = 0
            
            while self.crawl_queue and pages_processed < self.max_pages:
                current_url = self.crawl_queue.pop(0)
                
                msg = f"Processing ({pages_processed + 1}/{self.max_pages}): {current_url}"
                print(msg)
                if self.progress_callback:
                    self.progress_callback(pages_processed, self.max_pages, msg)
                
                try:
                    self.process_page(current_url)
                    pages_processed += 1
                except Exception as e:
                    print(f"Error processing {current_url}: {e}")
                
            print("Crawl Complete. Saving Metadata...")
            self.save_excel()
            
        except Exception as e:
            print(f"Critical Error: {e}")
        finally:
            try:
                self.driver.quit()
            except:
                pass

    def process_page(self, url):
        try:
            self.driver.get(url)
            time.sleep(2) # Wait for JS
        except TimeoutException:
            print(f"Timeout loading {url}")
            self.driver.execute_script("window.stop();")
        except Exception as e:
            print(f"Failed to load {url}: {e}")
            return

        soup = BeautifulSoup(self.driver.page_source, 'html.parser')
        
        # 1. Harvest Links for Queue
        self.harvest_links(soup, url)
        
        # 1.5 Handle PDFs found directly on page
        self.process_pdfs_on_page(soup, url)
        
        # 2. Extract Content
        title, body = self.extract_content(soup)
        
        if len(body) > 100:
            self.save_record(title, body, url)
        else:
            print(f"Skipping {url}: Content too short.")

    def process_pdfs_on_page(self, soup, current_url):
        for a in soup.find_all('a', href=True):
            href = a['href']
            full_url = urljoin(current_url, href)
            # Normalize URL (remove query params/anchors for extension check)
            pure_url = full_url.split('?')[0].split('#')[0]
            
            if pure_url.lower().endswith('.pdf'):
                if full_url not in self.visited_urls:
                    self.visited_urls.add(full_url)
                    pdf_name = a.get_text(strip=True) or "downloaded_document"
                    print(f"    [PDF Found] {full_url}")
                    
                    pdf_bytes = download_and_clean_pdf(full_url, pdf_name)
                    if pdf_bytes:
                        self.save_pdf_record(pdf_name, full_url, pdf_bytes)

    def save_pdf_record(self, name, url, content):
        safe_name = sanitize_filename(name) or "document"
        if not safe_name.lower().endswith('.pdf'):
            filename = f"{safe_name}.pdf"
        else:
            filename = safe_name
            
        save_path = os.path.join(self.output_dir, sanitize_filename(self.base_domain), "PDFs")
        if not os.path.exists(save_path): os.makedirs(save_path, exist_ok=True)
        
        full_path = os.path.join(save_path, filename)
        # Handle filename collisions if they occur (rare for same session)
        if os.path.exists(full_path):
            filename = f"{os.path.splitext(filename)[0]}_{uuid.uuid4().hex[:4]}.pdf"
            full_path = os.path.join(save_path, filename)
            
        with open(full_path, "wb") as f:
            f.write(content)
            
        rel_path = os.path.relpath(full_path, self.output_dir)
        self.metadata.append({
            "UUID": str(uuid.uuid4()),
            "Main_Category": self.base_domain,
            "Sub_Category": "PDF Document",
            "Entity_Title": name,
            "Description_Snippet": f"PDF Document from {url}",
            "Local_File_Path": rel_path,
            "TIPP_Source_URL": url,
            "External_Reference_URL": url
        })
        print(f"    Saved PDF: {filename}")

    def harvest_links(self, soup, current_url):
        for a in soup.find_all('a', href=True):
            href = a['href']
            full_url = urljoin(current_url, href)
            parsed = urlparse(full_url)
            
            # Simple Domain Check
            if parsed.netloc == self.base_domain:
                clean_url = full_url.split('#')[0]
                if clean_url not in self.visited_urls and clean_url not in self.crawl_queue:
                    if not any(ext in clean_url.lower() for ext in ['.jpg', '.png', '.pdf', '.zip', '.css', '.js']):
                        self.crawl_queue.append(clean_url)
                        self.visited_urls.add(clean_url)

    def extract_content(self, soup):
        clean_soup = soup
        for tag in clean_soup(['script', 'style', 'nav', 'footer', 'iframe', 'form', 'header']):
            tag.decompose()
            
        title = "Untitled Page"
        if soup.title:
            title = soup.title.get_text(strip=True)
            
        h1 = clean_soup.find('h1')
        if h1:
            title = h1.get_text(strip=True)
            
        content_body = ""
        main = clean_soup.find('main')
        if main:
            content_body = main.get_text(separator='\n', strip=True)
        else:
            article = clean_soup.find('article')
            if article:
                content_body = article.get_text(separator='\n', strip=True)
            else:
                content_body = clean_soup.get_text(separator='\n', strip=True)
                
        content_body = re.sub(r'\n\s*\n', '\n\n', content_body)
        return title, content_body.strip()

    def save_record(self, title, body, source_url):
        record_uuid = str(uuid.uuid4())
        safe_title = sanitize_filename(title) or "Untitled"
        filename = f"{safe_title[:100]}.docx"
        
        save_path = os.path.join(self.output_dir, sanitize_filename(self.base_domain))
        if not os.path.exists(save_path): os.makedirs(save_path, exist_ok=True)
            
        file_full_path = os.path.join(save_path, filename)
        if os.path.exists(file_full_path):
            filename = f"{os.path.splitext(filename)[0]}_{uuid.uuid4().hex[:4]}.docx"
            file_full_path = os.path.join(save_path, filename)
        
        try:
            doc = Document()
            p_head = doc.add_paragraph()
            run_head = p_head.add_run(f"Source: {self.base_domain}")
            run_head.font.size = Pt(10)
            run_head.font.color.rgb = RGBColor(128, 128, 128)
            
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = p_title.add_run(title)
            run_title.bold = True
            run_title.font.size = Pt(16)
            
            p_body = doc.add_paragraph(body)
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            section = doc.sections[0]
            footer = section.footer
            p_foot = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p_foot.text = "Source Link: "
            run_link = p_foot.add_run(source_url)
            run_link.font.color.rgb = RGBColor(0, 0, 255)
            
            doc.save(file_full_path)
            
            rel_path = os.path.relpath(file_full_path, self.output_dir)
            self.metadata.append({
                "UUID": record_uuid,
                "Main_Category": self.base_domain,
                "Sub_Category": "General",
                "Entity_Title": title,
                "Description_Snippet": body[:150],
                "Local_File_Path": rel_path,
                "TIPP_Source_URL": source_url,
                "External_Reference_URL": source_url
            })
            print(f"    Saved: {filename}")
        except Exception as e:
            print(f"Error saving docx: {e}")

    def save_excel(self):
        if not self.metadata: return
        try:
            df = pd.DataFrame(self.metadata)
            if os.path.exists(self.metadata_file):
                 try:
                     existing_df = pd.read_excel(self.metadata_file)
                     df = pd.concat([existing_df, df], ignore_index=True)
                 except: pass
            df.to_excel(self.metadata_file, index=False)
        except Exception as e:
            print(f"Error saving metadata: {e}")
