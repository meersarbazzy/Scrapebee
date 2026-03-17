import os
import time
import uuid
import re
import pandas as pd
from typing import List, Dict, Any
from urllib.parse import urljoin
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathvalidate import sanitize_filename
from scrapebee.core.pdf_processor import download_and_clean_pdf
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException

class TIPPScraperFinal:
    def __init__(self, output_dir=None, progress_callback=None):
        self.base_url = "https://tipp.gov.pk/"
        # Dynamic path management
        self.output_dir = output_dir or os.path.join(os.getcwd(), "data", "outputs")
        self.progress_callback = progress_callback
        self.metadata_file = os.path.join(self.output_dir, "metadata.xlsx")
        self.metadata = []
        self.visited_urls = set()
        self.crawl_queue = []
        
        # Setup Driver
        options = webdriver.ChromeOptions()
        options.add_argument("--start-maximized")
        options.add_argument("--page-load-strategy=eager") 
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        options.add_argument('--blink-settings=imagesEnabled=false')
        options.add_argument("--ignore-certificate-errors")
        options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        self.driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=options)
        
        # Anti-detection CDP
        self.driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
            "source": "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        })
        self.driver.set_page_load_timeout(120)
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir, exist_ok=True)

    def run(self):
        try:
            print("Phase 1: Crawling Menu Structure...")
            if self.progress_callback: self.progress_callback(0, 100, "Crawling Menu Structure...")
            menu_map = self.crawl_menu()
            print(f"Found {len(menu_map)} pages to scrape.")
            
            print("Phase 2: Processing Pages...")
            total = len(menu_map)
            for idx, item in enumerate(menu_map):
                msg = f"Processing {item['sub_cat']} ({idx+1}/{total})"
                print(f"[{idx+1}/{len(menu_map)}] {item['main_cat']} > {item['sub_cat']}")
                
                if self.progress_callback: 
                    self.progress_callback(idx+1, total, msg)
                
                if 'tradeInfo/searchProduct' in item['url']:
                    print(f"  [Auto-Mode] Skipping excluded page: {item['sub_cat']}")
                    continue

                self.process_page(item)
                
            print("Phase 3: Saving Metadata...")
            if self.progress_callback: self.progress_callback(total, total, "Saving Metadata...")
            self.save_excel()
            print("Done.")
            
        except Exception as e:
            print(f"Critical Error: {e}")
        finally:
            print("Finalizing: Saving metadata and quitting...")
            try: self.save_excel()
            except: pass
            try: self.driver.quit()
            except: pass

    def crawl_menu(self) -> List[Dict]:
        try:
            self.driver.get(self.base_url)
        except: pass
        
        time.sleep(5)
        try:
            src = self.driver.page_source
            soup = BeautifulSoup(src, 'html.parser')
        except: return []
        
        menu_items = []
        nav_ul = soup.find('ul', class_='navbar-nav')
        if not nav_ul: return []
            
        self._recurse_menu(nav_ul, [], menu_items)
        return menu_items

    def _recurse_menu(self, ul_element, path, result_list):
        children = ul_element.find_all('li', recursive=False)
        for li in children:
            a_tag = li.find('a', recursive=False) or li.find('a')
            if not a_tag: continue
                
            text = a_tag.get_text(strip=True)
            href = a_tag.get('href')
            new_path = path + [text]
            
            submenu = li.find('ul')
            if submenu:
                self._recurse_menu(submenu, new_path, result_list)
            else:
                if href and href != '#' and not href.startswith('javascript'):
                    full_url = urljoin(self.base_url, href)
                    if self.base_url in full_url:
                        main_cat = new_path[0] if len(new_path) > 0 else "Uncategorized"
                        sub_cat = new_path[1] if len(new_path) > 1 else "General"
                        result_list.append({
                            'url': full_url,
                            'path': new_path,
                            'main_cat': main_cat,
                            'sub_cat': sub_cat
                        })

    def process_page(self, item):
        url = item['url']
        path = item['path']
        current_url = url
        page_num = 1
        
        while True:
            try:
                self.driver.get(current_url)
            except TimeoutException:
                 try: self.driver.execute_script("window.stop();")
                 except: pass
            except: pass
            
            time.sleep(2)
            soup = BeautifulSoup(self.driver.page_source, 'html.parser')
            entities = self.extract_entities(soup, path)
            for ent in entities:
                self.save_record(ent['title'], ent['body'], ent['link'], item)
            
            self.process_pdfs_on_page(soup, current_url, item)
                
            next_url = None
            try:
                next_li = soup.find('li', class_='next')
                if next_li:
                    a_next = next_li.find('a', href=True)
                    if a_next: next_url = urljoin(self.base_url, a_next['href'])

                if not next_url:
                    a_tags = soup.find_all('a', href=True)
                    for a in a_tags:
                        t = a.get_text(strip=True).lower()
                        if 'next >' in t or 'next' == t or '»' in t:
                             pot_url = urljoin(self.base_url, a['href'])
                             if pot_url != current_url:
                                 next_url = pot_url
                                 break
            except: pass

            if next_url and next_url != current_url:
                 current_url = next_url
                 page_num += 1
            else: break

    def extract_entities(self, soup, path):
        entities = []
        main_content = soup.find('main') or soup.find(id='content') or soup.body
        if not main_content: return []
        
        for tag in main_content(['script', 'style', 'nav', 'footer', 'iframe', 'form']):
            tag.decompose()

        # Simplified Logic for brevity and health
        panels = main_content.find_all('div', class_=lambda c: c and 'panel' in c)
        if panels:
            for p in panels:
                heading = p.find('div', class_='panel-heading') or p.find(class_='panel-title')
                title = heading.get_text(strip=True) if heading else "Untitled"
                body_div = p.find('div', class_='panel-body') or p
                body_text = body_div.get_text(separator='\n', strip=True)
                
                link = None
                for a in p.find_all('a', href=True):
                    href = a['href']
                    if not href.startswith('#') and 'javascript' not in href:
                        link = href
                        break
                if len(body_text) > 50:
                    entities.append({'title': title, 'body': body_text.strip(), 'link': link})
            if entities: return entities

        # Fallback
        title = soup.title.get_text(strip=True) if soup.title else path[-1]
        body = main_content.get_text(separator='\n', strip=True)[:5000]
        entities.append({'title': title, 'body': body.strip(), 'link': urljoin(self.base_url, item['url'] if 'item' in locals() else '')})
        return entities

    def process_single_url(self, url):
        item = {'url': url, 'path': ['Manual', 'Batch'], 'main_cat': 'Batch', 'sub_cat': 'Single'}
        self.process_page(item)

    def process_pdfs_on_page(self, soup, current_url, item):
        for a in soup.find_all('a', href=True):
            href = a['href']
            if href.lower().endswith('.pdf'):
                pdf_url = urljoin(current_url, href)
                if not any(m['TIPP_Source_URL'] == pdf_url for m in self.metadata):
                    pdf_bytes = download_and_clean_pdf(pdf_url, "doc")
                    if pdf_bytes:
                        self.save_pdf_record(a.get_text(strip=True), pdf_url, pdf_bytes, item)

    def save_pdf_record(self, name, url, content, item):
        safe_name = sanitize_filename(name) or "document"
        if not safe_name.lower().endswith('.pdf'):
            filename = f"{safe_name}.pdf"
        else:
            filename = safe_name
            
        folder_names = [sanitize_filename(p) for p in item['path'][:-1]]
        save_path = os.path.join(self.output_dir, *folder_names, "PDFs")
        if not os.path.exists(save_path): os.makedirs(save_path, exist_ok=True)
        
        full_path = os.path.join(save_path, filename)
        if os.path.exists(full_path):
            filename = f"{os.path.splitext(filename)[0]}_{uuid.uuid4().hex[:4]}.pdf"
            full_path = os.path.join(save_path, filename)
            
        with open(full_path, "wb") as f: f.write(content)
        rel_path = os.path.relpath(full_path, self.output_dir)
        self.metadata.append({
            "UUID": str(uuid.uuid4()), "Main_Category": item['main_cat'],
            "Sub_Category": "PDF Document", "Entity_Title": name,
            "Description_Snippet": f"PDF from {url}", "Local_File_Path": rel_path,
            "TIPP_Source_URL": url, "External_Reference_URL": url
        })

    def save_record(self, title, body, ext_href, item):
        record_uuid = str(uuid.uuid4())
        safe_title = sanitize_filename(title) or "Untitled"
        filename = f"{safe_title[:100]}.docx"
        folder_names = [sanitize_filename(p) for p in item['path'][:-1]]
        save_path = os.path.join(self.output_dir, *folder_names)
        if not os.path.exists(save_path): os.makedirs(save_path, exist_ok=True)
        
        file_full_path = os.path.join(save_path, filename)
        if os.path.exists(file_full_path):
            filename = f"{os.path.splitext(filename)[0]}_{uuid.uuid4().hex[:4]}.docx"
            file_full_path = os.path.join(save_path, filename)
        try:
            doc = Document()
            doc.add_paragraph(f"{item['main_cat']} > {item['sub_cat']}")
            doc.add_paragraph(title)
            doc.add_paragraph(body)
            doc.save(file_full_path)
            rel_path = os.path.relpath(file_full_path, self.output_dir)
            self.metadata.append({
                "UUID": record_uuid, "Main_Category": item['main_cat'],
                "Sub_Category": item['sub_cat'], "Entity_Title": title,
                "Description_Snippet": body[:150], "Local_File_Path": rel_path,
                "TIPP_Source_URL": item['url'], "External_Reference_URL": ext_href
            })
        except: pass

    def save_excel(self):
        if not self.metadata: return
        try:
            df = pd.DataFrame(self.metadata)
            df.to_excel(self.metadata_file, index=False)
        except: pass

if __name__ == "__main__":
    TIPPScraperFinal().run()
