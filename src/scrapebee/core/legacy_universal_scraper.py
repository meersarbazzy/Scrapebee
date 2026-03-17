print("DEBUG: Scraper script started...", flush=True)
import os
try:
    import re
    import time
    from typing import List, Dict, Tuple
    from urllib.parse import urljoin

    from bs4 import BeautifulSoup
    from docx import Document
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.common.exceptions import StaleElementReferenceException, TimeoutException, NoSuchElementException, ElementClickInterceptedException
    from webdriver_manager.chrome import ChromeDriverManager
    from selenium.webdriver.chrome.service import Service as ChromeService
    from pathvalidate import sanitize_filename
    print("DEBUG: Imports successful.", flush=True)
except ImportError as e:
    print(f"CRITICAL IMPORT ERROR: {e}", flush=True)
    exit(1)

class UniversalScraper:
    def __init__(self, output_dir="c:/Project-DQ/Validata/Scrapper/universal_output"):
        self.base_url = "https://tipp.gov.pk/"
        self.output_dir = output_dir
        print("DEBUG: Setting up driver...", flush=True)
        self.driver = self._setup_driver()
        print("DEBUG: Driver setup complete.", flush=True)
        self.wait = WebDriverWait(self.driver, 10)
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def _setup_driver(self):
        options = webdriver.ChromeOptions()
        options.add_argument("--start-maximized")
        options.add_argument("--page-load-strategy=none")
        # options.add_argument("--headless") # Commented out for debugging
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        # Anti-detection settings
        options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
        options.add_argument("--ignore-certificate-errors")
        options.add_argument("--ignore-ssl-errors")
        
        # Robust Anti-Detection
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        options.add_argument('--disable-blink-features=AutomationControlled')
        
        service = ChromeService(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        # Execute CDP command to modify navigator.webdriver flag
        driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
            "source": """
                Object.defineProperty(navigator, 'webdriver', {
                    get: () => undefined
                })
            """
        })
        
        driver.set_page_load_timeout(60) # Increased timeout
        return driver

    def run(self):
        try:
            print(f"DEBUG: Navigating to {self.base_url}...", flush=True)
            self.driver.get(self.base_url)
            print("DEBUG: Navigation command sent. Waiting for body...", flush=True)
            
            # Explicit wait for body
            self.wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            # Optional: Wait for title to be non-empty
            # WebDriverWait(self.driver, 10).until(lambda d: d.title != "")
            
            # Give a bit of time for scripts to populate menu
            time.sleep(5) 
            
            print(f"DEBUG: Page Loaded. Title: {self.driver.title}", flush=True)
            self._handle_initial_popups()
            
            # Identify Main Navigation
            # Adjust selector based on actual site structure. 
            # Assuming a standard nav tag or specific ID often found in WP/Bootstrap sites
            # Specific to tipp.gov.pk, we need to find the main menu container.
            # I will assume a generic strategy first: finding the header nav.
            
            main_nav_items = self._find_main_menu_items()
            print(f"DEBUG: Found {len(main_nav_items)} main menu items.", flush=True)
            if not main_nav_items:
                 print("WARNING: No main menu items found. Check selectors.", flush=True)
            
            # Using index-based loop to handle stale elements
            for i in range(len(main_nav_items)):
                try:
                    # Re-find items at start of each iteration
                    current_nav_items = self._find_main_menu_items()
                    if i >= len(current_nav_items):
                        break
                        
                    item = current_nav_items[i]
                    item_text = item.text.strip()
                    if not item_text:
                        continue
                        
                    print(f"Processing Main Category: {item_text}")
                    
                    # Check if it has a dropdown
                    # Strategy: Click/Hover and check for submenu
                    # For now, we assume click triggers dropdown or visit
                    
                    self._process_menu_item(item, [item_text])
                    
                    # Return to base URL to reset state for next main category if needed
                    # Or rely on _process_menu_item to return to a safe state
                    self.driver.get(self.base_url)
                    
                except Exception as e:
                    print(f"Error processing main category index {i}: {e}")
                    self.driver.get(self.base_url)
                    
        finally:
            # self.driver.quit() # Keep open for debugging for now
            pass

    def _find_main_menu_items(self):
        # Heuristic selector for main menu - needs adjustment based on actual site
        # Often: //nav//ul/li or similar
        # Based on visual inspection of similar sites: #menu-main-menu > li, or .navbar-nav > li
        selectors = [
            "//ul[contains(@class, 'navbar-nav')]/li",
            "//div[contains(@class, 'menu')]//ul/li",
            "//nav//ul/li"
        ]
        
        for selector in selectors:
            elements = self.driver.find_elements(By.XPATH, selector)
            if elements and len(elements) > 3: # Reasonable number for a main menu
                return elements
        
        return []

    def _process_menu_item(self, element, path: List[str]):
        """
        Recursive function to handle menu items.
        element: The WebElement (li) to process
        path: List of strings representing the breadcrumb path
        """
        try:
            # check for submenu
            submenu = element.find_elements(By.XPATH, ".//ul")
            
            if submenu:
                # It has children
                print(f"  > Expanding {path[-1]}...")
                
                # Clicking parent to toggle dropdown might be needed
                # carefully check if it's a link or just a toggler
                link = element.find_element(By.TAG_NAME, "a")
                
                # Try hover or click
                webdriver.ActionChains(self.driver).move_to_element(link).perform()
                time.sleep(0.5) # Wait for animation
                
                # Re-find submenu text items to iterate
                # We need to find the direct children LIs of this submenu
                submenu_lis = submenu[0].find_elements(By.XPATH, "./li")
                
                # We can't iterate these elements directly if we navigate away.
                # So we store their basic info (like text or index) and re-find them.
                # Since menus can be complex, robust recursion is hard without DOM snapshot.
                # Simplified approach: Get links, if href is valid, visit.
                
                for j in range(len(submenu_lis)):
                    # Re-acquire parent and submenu context
                    # This part is the specific "re-find" logic to survive page loads if any
                    # But if we just hover, we might not navigate away yet?
                    # If clicking a child processes a page, we WILL navigate away.
                    
                    # Re-find strategy
                    parent = self._find_menu_item_by_text(path[-1]) # Helper needed?
                    if not parent:
                        break # Lost context
                        
                    current_submenu = parent.find_element(By.XPATH, ".//ul")
                    current_lis = current_submenu.find_elements(By.XPATH, "./li")
                    
                    if j >= len(current_lis):
                        break
                        
                    sub_li = current_lis[j]
                    sub_text = sub_li.text.split('\n')[0].strip() # Text might contain submenu text
                    if not sub_text:
                        sub_text = "Untitled"
                        
                    new_path = path + [sub_text]
                    
                    # Recurse or Visit
                    # If this sub_li has a UL, recurse. Else visit.
                    if sub_li.find_elements(By.XPATH, ".//ul"):
                        self._process_menu_item(sub_li, new_path)
                    else:
                        # Visit
                        link_tag = sub_li.find_element(By.TAG_NAME, "a")
                        href = link_tag.get_attribute("href")
                        if href and href != "#" and "javascript" not in href:
                             print(f"    -> Visiting: {sub_text} ({href})")
                             self._visit_and_scrape(href, new_path)
                             
                             # Restore menu state after return
                             # This is tricky. Easiest is to traverse from root again or back()
                             # _visit_and_scrape handles the navigation.
                             # After returning, the DOM is fresh. We need to re-open menus.
                             self._open_menu_path(path)
                             
            else:
                # No submenu, try to visit parent link
                link = element.find_element(By.TAG_NAME, "a")
                href = link.get_attribute("href")
                if href and href != "#":
                    print(f"  > Visiting Leaf: {path[-1]}")
                    self._visit_and_scrape(href, path)
                    
        except Exception as e:
            print(f"Error in _process_menu_item for {path}: {e}")

    def _open_menu_path(self, path):
        # Helper to re-expand menus to get back to state
        # Not fully implemented - requires robust selector logic matching text
        pass

    def _visit_and_scrape(self, url, path):
        self.driver.get(url)
        time.sleep(2) # Smart wait replacement
        
        while True: # Pagination loop
            self._scrape_current_page(path)
            
            # Check for next page
            if not self._go_to_next_page():
                break

    def _scrape_current_page(self, path):
        print(f"      Scraping content on {self.driver.current_url}")
        
        soup = BeautifulSoup(self.driver.page_source, 'html.parser')
        
        # Heuristic to find content blocks
        # 1. Look for typical article/content containers 
        # On tipp.gov.pk, content might be in generic divs
        
        content_found = False
        
        # Strategy: Look for blocks with H2/H3/H4 + P + A
        # Or look for specific known classes if available.
        # Fallback: Scrape the main content area.
        
        # Identify main content container
        main_content = soup.find('main') or soup.find(id='content') or soup.find(id='main') or soup.body
        
        if main_content:
            # Refined detection: List items or Repeated Blocks
            # Let's look for headings and following paragraphs
            blocks = []
            
            # Find all headings
            headings = main_content.find_all(['h1', 'h2', 'h3', 'h4'])
            
            for h in headings:
                title = h.get_text(strip=True)
                if not title:
                    continue
                
                # Gather content until next header
                body_parts = []
                links = []
                
                curr = h.next_sibling
                while curr:
                    if curr.name in ['h1', 'h2', 'h3', 'h4']:
                        break
                    
                    if curr.name: # Skip NavigableString if empty
                        text = curr.get_text(strip=True)
                        if text:
                            body_parts.append(text)
                        
                        # Find links
                        if curr.name == 'a':
                             links.append(curr.get('href'))
                        for a in curr.find_all('a'):
                             if a.get('href'):
                                links.append(a.get('href'))
                                
                    curr = curr.next_sibling
                
                if body_parts:
                    blocks.append({
                        'title': title,
                        'body': "\n".join(body_parts),
                        'links': links
                    })
            
            # Save blocks
            for block in blocks:
                content_found = True
                self._save_to_word(block, path)
        
        if not content_found:
             print("      [!] No structured content found. Dumping page text.")
             # Fallback: Dump valid text
             text = main_content.get_text(separator='\n', strip=True)
             self._save_to_word({'title': 'Page Content', 'body': text, 'links': []}, path)

    def _save_to_word(self, item, path):
        # Create folder path
        # [Main Category] / [Sub Category] ...
        # First item in path is Main, rest are Sub
        
        folder_structure = path[:-1] # Exclude leaf page name if it's the title?
        if not folder_structure:
            folder_structure = ["Uncategorized"]
            
        # Sanitize folder names
        safe_folders = [sanitize_filename(p) for p in folder_structure]
        save_dir = os.path.join(self.output_dir, *safe_folders)
        
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)
            
        # Filename
        title_slug = sanitize_filename(item['title'])[:50] # truncated
        filename = f"{title_slug}.docx"
        file_path = os.path.join(save_dir, filename)
        
        # Create Doc
        doc = Document()
        
        # Header: [Main Category] > [Sub Category] > [Page Title]
        header_text = " > ".join(path + [item['title']])
        p = doc.add_paragraph()
        runner = p.add_run(header_text)
        runner.bold = True
        
        # Title
        doc.add_heading(item['title'], 0)
        
        # Body
        doc.add_paragraph(item['body'])
        
        # Links
        if item['links']:
            doc.add_heading('References:', level=3)
            for link in item['links']:
                doc.add_paragraph(link, style='List Bullet')
                
        doc.save(file_path)
        print(f"      Saved: {file_path}")

    def _go_to_next_page(self):
        # Find next button
        try:
            # Common patterns for Next button
            next_btns = self.driver.find_elements(By.XPATH, "//a[contains(text(), 'Next') or contains(text(), '»')]")
            for btn in next_btns:
                if btn.is_displayed() and btn.is_enabled():
                    print("      Navigating to next page...")
                    btn.click()
                    time.sleep(2)
                    return True
        except:
            pass
        return False

    def _handle_initial_popups(self):
        # Stub for closing popups if any
        pass
        
    def _find_menu_item_by_text(self, text):
        """
        Re-finds a menu item (LI) based on its link text.
        Used to restore context after navigation.
        """
        try:
            # Try exact match on 'a' tag text inside 'li'
            # Note: normalize-space is important for clean text matching
            # XPath 1.0 doesn't support 'matches', so we use standard equality or contains
            
            # 1. Search in main menu candidates first?
            # Or just search global LI links? 
            # A global search might be too broad (footer links etc).
            # But the recursion context implies we know it's a menu item.
            
            # Using a fairly specific xpath looking for nav items
            xpath = f"//li[.//a[normalize-space()='{text}']]"
            elements = self.driver.find_elements(By.XPATH, xpath)
            
            # Filter for visible ones or ones inside 'nav' if possible
            for el in elements:
                if el.is_displayed():
                    return el
            
            if elements:
                return elements[0]
                
        except Exception as e:
            print(f"Error re-finding menu item '{text}': {e}")
            
        return None

if __name__ == "__main__":
    try:
        scraper = UniversalScraper()
        scraper.run()
    except Exception as e:
        import traceback
        print("CRITICAL ERROR IN MAIN:", flush=True)
        traceback.print_exc()
