import fitz  # PyMuPDF
import io
import os
import re
import uuid
import requests
import zipfile
import pytesseract
from PIL import Image
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch

def process_pdf_in_memory(input_bytes, filename_for_logging):
    """
    Cleans PDF by redacting header/footer zones (top 10% & bottom 10%).
    Copies pages intact to preserve fonts and Unicode text.
    """
    doc_in = None
    doc_out = None
    try:
        doc_in = fitz.open(stream=input_bytes, filetype="pdf")
        doc_out = fitz.open()

        for page_in in doc_in:
            # Copy page as-is into output doc (preserves fonts/Unicode)
            doc_out.insert_pdf(doc_in, from_page=page_in.number, to_page=page_in.number)
            page_out = doc_out[-1]

            h = page_out.rect.height
            w = page_out.rect.width

            # Define redaction zones (top 8% and bottom 8%)
            margin_pct = 0.08
            header_rect = fitz.Rect(0, 0, w, h * margin_pct)
            footer_rect = fitz.Rect(0, h * (1 - margin_pct), w, h)

            # Redact
            page_out.add_redact_annot(header_rect, fill=(1, 1, 1)) # White fill
            page_out.add_redact_annot(footer_rect, fill=(1, 1, 1))
            page_out.apply_redactions()

        return doc_out.tobytes()

    except Exception as e:
        print(f"Cleaning error for {filename_for_logging}: {e}")
        return None
    finally:
        if doc_in: doc_in.close()
        if doc_out: doc_out.close()

def create_text_pdf(text_content):
    """Creates a simple PDF with the provided text."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    text_object = c.beginText(inch, height - inch)
    text_object.setFont("Helvetica", 11)
    
    lines = text_content.split('\n')
    for line in lines:
        if text_object.getY() < inch:
            c.drawText(text_object)
            c.showPage()
            text_object = c.beginText(inch, height - inch)
            text_object.setFont("Helvetica", 11)
        text_object.textLine(line)
    
    c.drawText(text_object)
    c.save()
    buffer.seek(0)
    return buffer

def process_pdf_pages(reader, page_order, user_text):
    """Rearranges pages and appends new text."""
    writer = PdfWriter()
    messages = []
    
    if page_order:
        try:
            indices = [int(x.strip()) - 1 for x in page_order.split(',')]
            for idx in indices:
                if 0 <= idx < len(reader.pages):
                    writer.add_page(reader.pages[idx])
                else:
                    messages.append(f"Skipped page {idx+1} (out of range)")
        except:
            messages.append("Invalid page order format. Kept original order.")
            for page in reader.pages: writer.add_page(page)
    else:
        for page in reader.pages: writer.add_page(page)

    if user_text:
        text_pdf_buf = create_text_pdf(user_text)
        text_reader = PdfReader(text_pdf_buf)
        for page in text_reader.pages:
            writer.add_page(page)
            
    final_buffer = io.BytesIO()
    writer.write(final_buffer)
    final_buffer.seek(0)
    return final_buffer, messages

def pdf_to_word(pdf_bytes):
    """
    Converts a PDF to a Word (.docx) document.
    - Text-based PDFs: extracted directly via PyMuPDF.
    - Scanned/image PDFs: OCR'd via pytesseract with confidence filtering.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
        docx = DocxDocument()
        
        # Setup modern styling
        style = docx.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        ocr_used = False
        
        def _add_styled_line(doc, text, font_size=11, bold=False, color=(0,0,0)):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.bold = bold
            if color != (0,0,0):
                run.font.color.rgb = RGBColor(*color)
            return p

        def _clean_ocr_lines(lines):
            clean = []
            garbage_patterns = [
                r'^[^\w\s]{3,}$', # symbols like === or ---
                r'^[\\|/]{2,}$', # slashes like // or \\
                r'^[0-9\(\)\. ]{15,}$', # long gazette-style numbers
                r'^[A-Z]{1,2} [0-9]{3,}' # ID markers
            ]
            for ln in lines:
                text = ln.strip()
                if not text: continue
                if any(re.search(p, text) for p in garbage_patterns): continue
                if len(text) < 4 and not text.isalnum(): continue
                clean.append(text)
            return clean

        for page in doc_pdf:
            text = page.get_text().strip()
            
            if text:
                # Standard text-based PDF
                lines = text.split('\n')
                for ln in lines:
                    ln = ln.strip()
                    if ln:
                        # Simple heading detect: all caps short lines
                        is_heading = ln.isupper() and len(ln.split()) <= 8
                        _add_styled_line(docx, ln, font_size=14 if is_heading else 11, bold=is_heading)
            else:
                # No text found — render page as image and OCR it
                ocr_used = True
                mat = fitz.Matrix(3.0, 3.0) # 3x zoom ≈ 216 DPI — better OCR accuracy
                pix = page.get_pixmap(matrix=mat)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Get OCR data with confidence
                data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                
                current_line = []
                good_lines = []
                last_line_num = -1
                
                for i in range(len(data['text'])):
                    # Filter out purely non-alphanumeric junk with low confidence
                    word = data['text'][i].strip()
                    conf = data['conf'][i]
                    line_num = data['line_num'][i]
                    
                    if not word: continue
                    if conf < 40 and not word.isalnum(): continue # Drop low-conf garbage
                    
                    if line_num != last_line_num and current_line:
                        good_lines.append(" ".join(current_line))
                        current_line = []
                    
                    current_line.append(word)
                    last_line_num = line_num
                
                if current_line:
                    good_lines.append(" ".join(current_line))

                for ln in _clean_ocr_lines(good_lines):
                    is_heading_candidate = ln.isupper() and len(ln.split()) <= 8
                    _add_styled_line(docx, ln, font_size=13 if is_heading_candidate else 11, 
                                     bold=is_heading_candidate)

        doc_pdf.close()
        buffer = io.BytesIO()
        docx.save(buffer)
        buffer.seek(0)
        return buffer, ocr_used

    except Exception as e:
        print(f"PDF to Word error: {e}")
        return None, False

def pdf_to_images(pdf_bytes):
    """Converts PDF pages to a ZIP of image files."""
    import fitz
    import zipfile
    import io
    
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_f:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # High resolution
            img_data = pix.tobytes("png")
            zip_f.writestr(f"page_{i+1}.png", img_data)
    doc.close()
    zip_buffer.seek(0)
    return zip_buffer

def pdf_to_excel(pdf_bytes):
    """Extracts text from PDF and saves it as a structured Excel spreadsheet."""
    import pandas as pd
    import fitz
    import io
    
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    data = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines:
            data.append({"Page": i+1, "Content": line})
    doc.close()
    
    df = pd.DataFrame(data)
    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="PDF_Content")
    excel_buffer.seek(0)
    return excel_buffer

def word_to_pdf(docx_bytes):
    """
    Converts a Word (.docx) file to PDF.
    Uses python-docx for extraction and ReportLab for high-fidelity rendering.
    """
    from docx import Document
    import io
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    docx_io = io.BytesIO(docx_bytes)
    doc = Document(docx_io)
    
    pdf_buffer = io.BytesIO()
    doc_pdf = SimpleDocTemplate(pdf_buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    for para in doc.paragraphs:
        if para.text.strip():
            # Basic style mapping
            style = styles['Normal']
            if para.style.name.startswith('Heading'):
                style = styles[para.style.name] if para.style.name in styles else styles['Heading1']
            
            p = Paragraph(para.text, style)
            story.append(p)
            story.append(Spacer(1, 0.1*inch))

    doc_pdf.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer

def excel_to_pdf(xlsx_bytes):
    """
    Converts Excel (.xlsx) data to a PDF table.
    """
    import pandas as pd
    import io
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    df = pd.read_excel(io.BytesIO(xlsx_bytes))
    pdf_buffer = io.BytesIO()
    
    # Use landscape if table is wide
    doc = SimpleDocTemplate(pdf_buffer, pagesize=landscape(A4) if len(df.columns) > 5 else A4)
    
    data = [df.columns.to_list()] + df.values.tolist()
    
    # Simple table styling
    t = Table(data)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    doc.build([t])
    pdf_buffer.seek(0)
    return pdf_buffer

def images_to_pdf(list_of_image_bytes):
    """
    Stitches multiple images into a single PDF document.
    """
    from PIL import Image
    import io
    
    pdf_buffer = io.BytesIO()
    img_list = []
    
    for img_bytes in list_of_image_bytes:
        img = Image.open(io.BytesIO(img_bytes)).convert('RGB')
        img_list.append(img)
    
    if img_list:
        img_list[0].save(pdf_buffer, format='PDF', save_all=True, append_images=img_list[1:])
    
    pdf_buffer.seek(0)
    return pdf_buffer

def compress_file(file_bytes, extension):
    """
    Applies specialized compression based on the file type.
    """
    import io
    import zipfile
    
    ext = extension.lower().strip('.')
    
    if ext == 'pdf':
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        buffer = io.BytesIO()
        # garbage=4: remove unused objects, clean up, etc.
        # deflate=True: compress streams
        doc.save(buffer, garbage=4, deflate=True, clean=True)
        doc.close()
        buffer.seek(0)
        return buffer

    elif ext in ['docx', 'xlsx']:
        # Office files are ZIPs. Re-compress with highest level.
        in_zip = io.BytesIO(file_bytes)
        out_zip = io.BytesIO()
        with zipfile.ZipFile(in_zip, 'r') as zin:
            with zipfile.ZipFile(out_zip, 'w', compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zout:
                for item in zin.infolist():
                    zout.writestr(item, zin.read(item.filename))
        out_zip.seek(0)
        return out_zip

    elif ext == 'csv':
        # CSV can be compressed into a GZIP stream
        import gzip
        buffer = io.BytesIO()
        with gzip.GzipFile(fileobj=buffer, mode='wb', compresslevel=9) as f:
            f.write(file_bytes)
        buffer.seek(0)
        return buffer

    elif ext == 'parquet':
        # Parquet optimization: use ZSTD high-level compression
        import pandas as pd
        import io
        df = pd.read_parquet(io.BytesIO(file_bytes))
        buffer = io.BytesIO()
        df.to_parquet(buffer, compression='zstd', index=False)
        buffer.seek(0)
        return buffer
    
    return io.BytesIO(file_bytes) # Default: no change

def download_and_clean_pdf(url, name_hint="document"):
    """Helper for main scrapers to handle PDFs on the fly."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'application/pdf,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8'
        }
        r = requests.get(url, headers=headers, timeout=30)
        r.raise_for_status()
        cleaned = process_pdf_in_memory(r.content, name_hint)
        return cleaned or r.content
    except:
        return None

class UniversalDocumentExtractor:
    """Discovers and downloads PDFs from a domain."""
    def __init__(self, start_url, max_pages=20, direct_only=False, progress_callback=None):
        self.start_url = start_url
        self.domain = urlparse(start_url).netloc
        self.max_pages = max_pages
        self.direct_only = direct_only
        self.progress_callback = progress_callback
        self.visited = set()
        self.to_visit = [start_url]
        self.discovered_pdfs = {} # url -> bytes
        self.metadata = [] # List of dicts: {source_url, title, filename, youtube_links}

    def run(self):
        count = 0
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8'
        }
        while self.to_visit and count < self.max_pages:
            url = self.to_visit.pop(0)
            if url in self.visited: continue
            self.visited.add(url)
            count += 1
            
            if self.progress_callback:
                self.progress_callback(count, self.max_pages, f"Scanning: {url}")
            
            try:
                r = requests.get(url, timeout=15, headers=headers)
                content_type = r.headers.get('Content-Type', '').lower()
                
                if 'application/pdf' in content_type:
                    self.discovered_pdfs[url] = r.content
                    continue
                
                if self.direct_only and count > 1: continue
                
                soup = BeautifulSoup(r.text, 'html.parser')
                
                # Capture YouTube Links
                page_youtube = []
                for link in soup.find_all('a', href=True):
                    href = link['href']
                    if 'youtube.com' in href or 'youtu.be' in href:
                        page_youtube.append(href)
                
                page_title = soup.title.string if soup.title else url
                
                for a in soup.find_all('a', href=True):
                    href = a['href']
                    full_url = urljoin(url, href)
                    
                    # Normalize URL (remove query params for extension check)
                    pure_url = full_url.split('?')[0].split('#')[0]
                    
                    if pure_url.lower().endswith('.pdf'):
                        if full_url not in self.discovered_pdfs:
                            try:
                                pdf_r = requests.get(full_url, timeout=15, headers=headers)
                                if 'application/pdf' in pdf_r.headers.get('Content-Type', '').lower():
                                    self.discovered_pdfs[full_url] = pdf_r.content
                                    filename = os.path.basename(pure_url)
                                    self.metadata.append({
                                        "Source Link": url,
                                        "Page Title": page_title,
                                        "Local File Path": filename,
                                        "YouTube Links": ", ".join(list(set(page_youtube))),
                                        "Scrape Date": time.strftime("%Y-%m-%d %H:%M:%S")
                                    })
                                    print(f"    [PDF Found] {full_url}")
                            except: pass
                    elif urlparse(full_url).netloc == self.domain:
                        if full_url not in self.visited and full_url not in self.to_visit:
                            # Avoid known non-html extensions
                            if not any(pure_url.lower().endswith(ext) for ext in ['.jpg', '.png', '.zip', '.docx', '.xlsx']):
                                self.to_visit.append(full_url)
            except Exception as e:
                print(f"Error scanning {url}: {e}")
