from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from docx import Document
from pathvalidate import sanitize_filename
import os
import time

def fetch_page(url):
    """Fetches the content of a URL using Selenium WebDriver."""
    driver = None
    try:
        options = Options()
        options.add_argument("--headless=new") # Run in headless mode
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        options.add_argument("--window-size=1920,1080")
        options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
        
        # Suppress log messages
        options.add_argument("--log-level=3")
        options.page_load_strategy = 'eager' # Don't wait for all resources (images/css) to finish

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        driver.set_page_load_timeout(60)
        driver.get(url)
        
        # Give it a moment to load dynamic content if any
        time.sleep(3)
        
        html = driver.page_source
        return html, None
        
    except Exception as e:
        print(f"Error fetching {url}: {e}")
        return None, str(e)
    finally:
        if driver:
            driver.quit()

def extract_content(html, url):
    """Extracts structured content from HTML."""
    soup = BeautifulSoup(html, 'html.parser')
    
    # Extract Title
    # Default to page title
    page_title = soup.title.string.strip() if soup.title else "Untitled Page"
    
    # Try to find a better content title from headings
    # We prioritize H1 -> H2 -> H4 (Found 'About TIPP' in H4) -> H3 (Disclaimer found here)
    # Also added check for bold/strong/span elements that might act as titles if headings are missing
    content_title = None
    
    # Priority list of tags to check for title
    title_tags = ['h1', 'h2', 'h4', 'strong', 'b', 'span']
    
    for tag in title_tags:
        found_tags = soup.find_all(tag)
        for t in found_tags:
            text = t.get_text(strip=True)
            # Filter out common noise
            # Title should be reasonably short (< 100 chars) but not too short (> 3 chars)
            if text and len(text) > 3 and len(text) < 100 and "disclaimer" not in text.lower() and "menu" not in text.lower():
                
                # Heuristic: If it's a span/b/strong, it might be just body text. 
                # Check if it has a specific class or seems isolated?
                # For this specific site, 'Forward Transit' was in a span.
                # Let's verify if it matches breadcrumbs or active menu?
                # A simple heuristic: first significant text that looks like a title.
                content_title = text
                break
        if content_title:
            break
            
    final_title = content_title if content_title else page_title
    
    content_elements = []
    
    # Clean up
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
        tag.decompose()
        
    body = soup.body
    if not body:
        return {'title': title, 'elements': [], 'url': url}

    tags_to_extract = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'div', 'span', 'article', 'section']
    # Expanded tags to ensure we catch content in divs/spans if p is missing
    # But we need to filter out noise. 
    # Let's stick to the previous list but maybe be more recursive or smart?
    # For now, stick to the previous list plus 'div' if it contains significant text directly.
    
    # Refined list
    tags_to_extract = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']

    for element in body.find_all(tags_to_extract):
        text = element.get_text(strip=True)
        if not text:
            continue
            
        # Avoid duplicate text (nested elements)
        # If a p is inside a div, we might catch both if we used div. 
        # Since we use p, h, li, we should be mostly fine, but nested lists might be duplicated if we aren't careful.
        # bs4 find_all is recursive by default.
        
        links = []
        if element.name == 'p' or element.name == 'li':
            for a in element.find_all('a', href=True):
                link_text = a.get_text(strip=True)
                if link_text:
                    links.append({'text': link_text, 'href': a['href']})

        content_elements.append({
            'type': element.name,
            'text': text,
            'links': links
        })
        
    return {
        'title': final_title,
        'elements': content_elements,
        'url': url
    }

def save_to_docx(content_data, output_dir="output"):
    """Saves the extracted content to a Word document."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    sanitized_title = sanitize_filename(content_data['title'])
    if not sanitized_title:
        sanitized_title = "scraped_page"
        
    filename = f"{sanitized_title}.docx"
    filepath = os.path.join(output_dir, filename)
    
    # Handle duplicates
    if os.path.exists(filepath):
        timestamp = int(time.time())
        filename = f"{sanitized_title}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)

    doc = Document()
    doc.add_heading(content_data['title'], 0)
    doc.add_paragraph(f"Source: {content_data['url']}")
    
    for item in content_data['elements']:
        text = item['text']
        
        if item['type'] == 'h1':
            doc.add_heading(text, level=1)
        elif item['type'] == 'h2':
            doc.add_heading(text, level=2)
        elif item['type'] == 'h3':
            doc.add_heading(text, level=3)
        elif item['type'] in ['h4', 'h5', 'h6']:
            doc.add_heading(text, level=4)
        elif item['type'] in ['p', 'li']:
            p = doc.add_paragraph()
            if item['type'] == 'li':
                p.style = 'List Bullet' 
            
            p.add_run(text)
            
            if item['links']:
                p.add_run(" (Links: ")
                for i, link in enumerate(item['links']):
                    p.add_run(f"{link['text']} -> {link['href']}")
                    if i < len(item['links']) - 1:
                        p.add_run(", ")
                p.add_run(")")

    doc.save(filepath)
    return filepath

def crawl_domain(start_url, max_pages=50):
    """Recursively crawls a domain starting from start_url."""
    
    from urllib.parse import urlparse, urljoin
    
    # Setup
    visited = set()
    queue = [start_url]
    domain = urlparse(start_url).netloc
    results = []
    
    # Selenium Driver (Reused for efficiency if possible, or new per page to be safe)
    # For efficiency we should ideally keep driver open, but our fetch_page opens/closes it.
    # To avoid 50 browser opens, let's refactor fetch_page OR just accept the overhead for now (safer against memory leaks).
    # Given the user wants "robust", restart is safer but slower.
    
    pages_crawled = 0
    
    while queue and pages_crawled < max_pages:
        url = queue.pop(0)
        
        if url in visited:
            continue
            
        visited.add(url)
        pages_crawled += 1
        
        # Fetch
        print(f"Crawling: {url}")
        html, error = fetch_page(url)
        
        if html:
            # Extract
            content_data = extract_content(html, url)
            results.append(content_data)
            
            # Find new links
            soup = BeautifulSoup(html, 'html.parser')
            for a in soup.find_all('a', href=True):
                href = a['href']
                full_url = urljoin(url, href)
                parsed = urlparse(full_url)
                
                # Internal links only
                if parsed.netloc == domain or parsed.netloc == '':
                    # Filter files and junk
                    clean_url = full_url.split('#')[0]
                    if clean_url not in visited and clean_url not in queue:
                        # rigorous filter
                        if not any(clean_url.endswith(ext) for ext in ['.pdf', '.jpg', '.png', '.css', '.js', '.zip', '.docx']):
                             queue.append(clean_url)
        else:
             print(f"Failed to crawl {url}: {error}")
             
    return results
